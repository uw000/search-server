from pathlib import Path

from docx import Document

from app.parsers.base import BaseParser, ParseResult
from app.parsers.text_cleaner import clean_text


class DocxParser(BaseParser):
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    def parse(self, file_path: Path) -> ParseResult:
        result = ParseResult()

        size_error = self.check_file_size(file_path)
        if size_error:
            result.errors.append(size_error)
            return result

        try:
            doc = Document(str(file_path))
        except Exception as e:
            result.errors.append(f"Failed to open DOCX: {e}")
            return result

        result.title = doc.core_properties.title or file_path.stem
        result.author = doc.core_properties.author

        sections: list[tuple[str | None, str]] = []
        current_heading: str | None = None
        current_text: list[str] = []

        for para in doc.paragraphs:
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                if current_text:
                    sections.append((current_heading, "\n".join(current_text)))
                    current_text = []
                current_heading = para.text
            else:
                if para.text.strip():
                    current_text.append(para.text)

        if current_text:
            sections.append((current_heading, "\n".join(current_text)))

        section_num = 0
        for heading, text in sections:
            text = clean_text(text, drop_page_numbers=False)
            if not text.strip():
                continue

            section_num += 1
            result.chunks.extend(
                self.chunk_long_text(
                    text,
                    page_number=section_num,
                    section=heading,
                    content_type="text",
                )
            )

        # 표(Table) 추출 — 본문 순회와 별개 경로.
        # 각 표를 하나의 "table" 타입 청크로 저장한다. 행은 "\n"로, 셀은 " | "로 구분.
        # 중첩 셀 내부의 표는 python-docx 가 cell.tables 로 노출하지만, 단순화를 위해
        # cell.text 의 평탄화된 텍스트만 사용한다(중첩 표는 부모 셀 텍스트에 포함됨).
        for t_idx, table in enumerate(doc.tables, start=1):
            rows_text: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows_text.append(" | ".join(cells))
            table_text = "\n".join(rows_text).strip()
            if not table_text:
                continue
            section_num += 1
            result.chunks.extend(
                self.chunk_long_text(
                    table_text,
                    page_number=section_num,
                    section=f"Table {t_idx}",
                    content_type="table",
                )
            )

        result.chunks = self.merge_small_chunks(result.chunks)
        result.total_pages = section_num
        return result
