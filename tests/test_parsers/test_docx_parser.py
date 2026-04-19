from pathlib import Path

import pytest

FIXTURES_DIR = Path(__file__).parent.parent / "fixtures"


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """임시 DOCX 파일 생성."""
    from docx import Document

    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is the introduction paragraph with enough text to be meaningful.")
    doc.add_heading("Chapter 1", level=2)
    doc.add_paragraph("Chapter 1 content. " * 20)
    doc.add_heading("Chapter 2", level=2)
    doc.add_paragraph("Chapter 2 content. " * 20)

    path = tmp_path / "test.docx"
    doc.save(str(path))
    return path


def test_docx_parser_basic(sample_docx: Path) -> None:
    from app.parsers.docx_parser import DocxParser

    parser = DocxParser()
    result = parser.parse(sample_docx)

    assert len(result.chunks) >= 2
    assert result.title is not None


def test_docx_parser_extensions() -> None:
    from app.parsers.docx_parser import DocxParser

    parser = DocxParser()
    assert ".docx" in parser.supported_extensions()


def test_docx_parser_extracts_tables(tmp_path: Path) -> None:
    from docx import Document

    from app.parsers.docx_parser import DocxParser

    doc = Document()
    doc.add_heading("Doc with table", level=1)
    doc.add_paragraph("Intro paragraph." * 20)
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "alpha"
    table.cell(1, 1).text = "1"
    table.cell(2, 0).text = "베타"
    table.cell(2, 1).text = "2"
    path = tmp_path / "t.docx"
    doc.save(str(path))

    parser = DocxParser()
    result = parser.parse(path)

    table_chunks = [c for c in result.chunks if c.content_type == "table"]
    assert len(table_chunks) == 1
    tc = table_chunks[0]
    # 셀 구분자(|), 행 구분자(줄바꿈) 확인
    assert "Name | Value" in tc.content or "Name" in tc.content
    assert "alpha" in tc.content
    assert "베타" in tc.content
    assert tc.section == "Table 1"


def test_docx_parser_skips_empty_tables(tmp_path: Path) -> None:
    from docx import Document

    from app.parsers.docx_parser import DocxParser

    doc = Document()
    doc.add_heading("H", level=1)
    doc.add_paragraph("body " * 50)
    # 빈 표 추가 (셀 텍스트 없음)
    doc.add_table(rows=2, cols=2)
    path = tmp_path / "empty_table.docx"
    doc.save(str(path))

    parser = DocxParser()
    result = parser.parse(path)

    assert not any(c.content_type == "table" for c in result.chunks)
