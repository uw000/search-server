"""테스트용 샘플 문서 파일을 생성하는 스크립트.

사용법:
    python -m tests.fixtures.create_samples
    python -m tests.fixtures.create_samples --output-dir /data/documents
"""

import argparse
from pathlib import Path


def create_sample_txt(output_dir: Path) -> Path:
    path = output_dir / "sample.txt"
    path.write_text(
        "Python 프로그래밍 가이드\n\n"
        "Chapter 1: 리스트와 튜플\n\n"
        "Python의 list comprehension은 매우 강력한 기능입니다. "
        "리스트 컴프리헨션을 사용하면 간결하고 읽기 쉬운 코드를 작성할 수 있습니다.\n\n"
        "squares = [x**2 for x in range(10)]\n\n"
        "이 코드는 0부터 9까지의 제곱수를 생성합니다.\n\n"
        "Chapter 2: 딕셔너리\n\n"
        "딕셔너리는 키-값 쌍을 저장하는 자료구조입니다. "
        "해시 테이블을 기반으로 구현되어 있어 O(1) 시간에 조회가 가능합니다.\n\n"
        "data = {'name': 'Alice', 'age': 30}\n\n"
        "Chapter 3: 함수형 프로그래밍\n\n"
        "map, filter, reduce 같은 함수형 프로그래밍 패턴을 Python에서도 사용할 수 있습니다. "
        "lambda 표현식과 함께 사용하면 더욱 강력합니다.\n\n"
        "result = list(map(lambda x: x * 2, [1, 2, 3, 4, 5]))\n",
        encoding="utf-8",
    )
    return path


def create_sample_pdf(output_dir: Path) -> Path:
    import fitz

    path = output_dir / "sample_text.pdf"
    doc = fitz.open()

    pages = [
        ("Chapter 1: 데이터베이스 기초",
         "관계형 데이터베이스는 테이블 형태로 데이터를 저장합니다.\n"
         "SQL(Structured Query Language)을 사용하여 데이터를 조회하고 조작할 수 있습니다.\n\n"
         "SELECT * FROM users WHERE age > 20;\n\n"
         "이 쿼리는 나이가 20세 이상인 모든 사용자를 조회합니다.\n"
         "JOIN 연산을 통해 여러 테이블의 데이터를 결합할 수 있습니다."),
        ("Chapter 2: 인덱스와 성능",
         "데이터베이스 인덱스는 B-Tree 구조를 사용하여 빠른 검색을 지원합니다.\n"
         "적절한 인덱스 설계는 쿼리 성능을 크게 향상시킬 수 있습니다.\n\n"
         "CREATE INDEX idx_users_email ON users(email);\n\n"
         "인덱스는 읽기 성능을 높이지만 쓰기 성능에는 약간의 오버헤드가 있습니다.\n"
         "복합 인덱스(composite index)를 사용하면 여러 컬럼에 대한 검색을 최적화할 수 있습니다."),
        ("Chapter 3: 트랜잭션",
         "ACID 속성은 데이터베이스 트랜잭션의 안정성을 보장합니다.\n"
         "Atomicity(원자성), Consistency(일관성), Isolation(격리성), Durability(내구성).\n\n"
         "BEGIN TRANSACTION;\n"
         "UPDATE accounts SET balance = balance - 1000 WHERE id = 1;\n"
         "UPDATE accounts SET balance = balance + 1000 WHERE id = 2;\n"
         "COMMIT;\n\n"
         "이 트랜잭션은 계좌 이체를 원자적으로 수행합니다."),
    ]

    for title, content in pages:
        page = doc.new_page()
        page.insert_text(fitz.Point(72, 72), title, fontsize=16)
        page.insert_text(fitz.Point(72, 110), content, fontsize=11)

    doc.ez_save(str(path))
    doc.close()
    return path


def create_sample_docx(output_dir: Path) -> Path:
    from docx import Document

    path = output_dir / "sample.docx"
    doc = Document()

    doc.add_heading("네트워크 프로토콜 개요", level=1)

    doc.add_heading("TCP/IP 모델", level=2)
    doc.add_paragraph(
        "TCP/IP는 인터넷의 기본 통신 프로토콜 스택입니다. "
        "Application Layer, Transport Layer, Internet Layer, Network Access Layer의 "
        "4개 계층으로 구성됩니다. TCP는 연결 지향적이며 신뢰성 있는 데이터 전송을 보장합니다."
    )

    doc.add_heading("HTTP 프로토콜", level=2)
    doc.add_paragraph(
        "HTTP(HyperText Transfer Protocol)는 웹에서 데이터를 주고받기 위한 프로토콜입니다. "
        "요청(Request)과 응답(Response) 구조로 동작하며, "
        "GET, POST, PUT, DELETE 등의 메서드를 지원합니다. "
        "HTTP/2는 멀티플렉싱을 통해 성능을 크게 향상시켰습니다."
    )

    doc.add_heading("DNS 시스템", level=2)
    doc.add_paragraph(
        "DNS(Domain Name System)은 도메인 이름을 IP 주소로 변환하는 분산 데이터베이스 시스템입니다. "
        "A 레코드, CNAME, MX 레코드 등 다양한 레코드 타입을 지원합니다. "
        "DNS 캐싱은 반복적인 조회를 줄여 응답 시간을 단축합니다."
    )

    doc.save(str(path))
    return path


def create_sample_epub(output_dir: Path) -> Path:
    from ebooklib import epub

    path = output_dir / "sample.epub"
    book = epub.EpubBook()

    book.set_identifier("sample-book-001")
    book.set_title("머신러닝 입문")
    book.set_language("ko")
    book.add_author("테스트 저자")

    ch1 = epub.EpubHtml(title="지도학습", file_name="ch1.xhtml", lang="ko")
    ch1.content = (
        "<h1>Chapter 1: 지도학습</h1>"
        "<p>지도학습(Supervised Learning)은 레이블이 있는 훈련 데이터를 사용합니다. "
        "분류(Classification)와 회귀(Regression)로 나뉩니다.</p>"
        "<p>선형 회귀, 로지스틱 회귀, 결정 트리, 랜덤 포레스트, SVM 등이 대표적인 알고리즘입니다.</p>"
    )

    ch2 = epub.EpubHtml(title="비지도학습", file_name="ch2.xhtml", lang="ko")
    ch2.content = (
        "<h1>Chapter 2: 비지도학습</h1>"
        "<p>비지도학습(Unsupervised Learning)은 레이블이 없는 데이터에서 패턴을 찾습니다. "
        "클러스터링(Clustering)과 차원 축소(Dimensionality Reduction)가 주요 기법입니다.</p>"
        "<p>K-Means, DBSCAN, PCA, t-SNE 등의 알고리즘이 사용됩니다.</p>"
    )

    ch3 = epub.EpubHtml(title="딥러닝", file_name="ch3.xhtml", lang="ko")
    ch3.content = (
        "<h1>Chapter 3: 딥러닝</h1>"
        "<p>딥러닝(Deep Learning)은 인공 신경망을 기반으로 한 머신러닝 기법입니다. "
        "CNN은 이미지 인식에, RNN/LSTM은 시계열 데이터에, Transformer는 자연어 처리에 주로 사용됩니다.</p>"
        "<p>역전파(Backpropagation) 알고리즘으로 가중치를 학습합니다.</p>"
    )

    book.add_item(ch1)
    book.add_item(ch2)
    book.add_item(ch3)

    book.toc = [
        epub.Link("ch1.xhtml", "지도학습", "ch1"),
        epub.Link("ch2.xhtml", "비지도학습", "ch2"),
        epub.Link("ch3.xhtml", "딥러닝", "ch3"),
    ]

    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", ch1, ch2, ch3]

    epub.write_epub(str(path), book)
    return path


def main() -> None:
    parser = argparse.ArgumentParser(description="Create sample document files for testing")
    parser.add_argument("--output-dir", type=str, default="tests/fixtures")
    args = parser.parse_args()

    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"Creating sample documents in {output_dir}...")

    txt = create_sample_txt(output_dir)
    print(f"  Created: {txt}")

    pdf = create_sample_pdf(output_dir)
    print(f"  Created: {pdf}")

    docx = create_sample_docx(output_dir)
    print(f"  Created: {docx}")

    epub = create_sample_epub(output_dir)
    print(f"  Created: {epub}")

    print("Done!")


if __name__ == "__main__":
    main()
